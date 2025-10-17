from __future__ import annotations

import json
import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

try:
    import seaborn as sns  # type: ignore
except ImportError:  # pragma: no cover
    sns = None

from docx import Document
from docx.shared import Inches

from .optimizer import OptimizationResult
from .settings import PipelineConfig

LOGGER = logging.getLogger("next_pipeline.reporting")


def create_heatmap(
    data: pd.DataFrame,
    output_path: Path,
    title: str = "",
    cmap: str = "RdYlGn_r",
) -> Path:
    """Generate a heatmap image from the provided DataFrame."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    plt.figure(figsize=(max(6, data.shape[1] * 0.35), max(6, data.shape[0] * 0.35)))
    if sns is not None:
        ax = sns.heatmap(
            data,
            cmap=cmap,
            center=0,
            linewidths=0.1,
            linecolor="white",
            cbar_kws={"label": "Z-Score"},
        )
    else:  # pragma: no cover - seaborn missing
        ax = plt.imshow(data, cmap=cmap, aspect="auto", interpolation="nearest")
        plt.colorbar(ax, label="Z-Score")
        ax = plt.gca()
    plt.title(title or "Signal Heatmap")
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()
    LOGGER.info("Saved heatmap to %s", output_path)
    return output_path


@dataclass
class ReportOrchestrator:
    config: PipelineConfig

    def summary_dataframe(
        self, results: Dict[str, OptimizationResult]
    ) -> pd.DataFrame:
        records = []
        for code, result in results.items():
            metrics = result.best_trial.test_metrics
            records.append(
                {
                    "code": code,
                    "rebalance_freq": result.best_params["rebalance_freq"],
                    "min_periods": result.best_params["min_periods"],
                    "future_days": result.best_params["future_days"],
                    "ic_lookback_period": result.best_params["ic_lookback_period"],
                    "test_sharpe": metrics["sharpe"],
                    "annual_return": metrics["annual_return"],
                    "annual_volatility": metrics["annual_volatility"],
                    "max_drawdown": metrics["max_drawdown"],
                    "calmar": metrics["calmar"],
                    "win_rate": metrics["win_rate"],
                    "profit_ratio": metrics["profit_ratio"],
                }
            )
        df = pd.DataFrame(records)
        if df.empty:
            return df
        df = df.sort_values(by="test_sharpe", ascending=False).reset_index(drop=True)
        return df

    def save_summary(
        self, summary_df: pd.DataFrame, output_dir: Optional[Path] = None
    ) -> Path:
        output_dir = output_dir or self.config.optimisation_output_dir()
        csv_name = self.config.reporting.get(
            "summary_filename", "指数参数绩效汇总.csv"
        )
        path = output_dir / csv_name
        summary_df.to_csv(path, index=False)
        LOGGER.info("Saved summary CSV to %s", path)
        return path

    def build_word_report(
        self,
        summary_df: pd.DataFrame,
        results: Dict[str, OptimizationResult],
        heatmap_path: Optional[Path] = None,
        output_dir: Optional[Path] = None,
    ) -> Path:
        output_dir = output_dir or self.config.optimisation_output_dir()
        word_name = self.config.reporting.get(
            "word_filename", "行业均值方差择时报告.docx"
        )
        doc = Document()
        doc.add_heading("行业均值方差择时策略报告", 0)

        doc.add_paragraph(
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        doc.add_paragraph(
            f"覆盖指数数：{summary_df.shape[0]}，数据起始："
            f"{self.config.data.get('start_date', 'N/A')}，结束："
            f"{self.config.data.get('end_date') or '最新'}。"
        )

        doc.add_heading("1. 参数概览", level=1)
        if summary_df.empty:
            doc.add_paragraph("无可用结果。")
        else:
            columns = [
                "code",
                "rebalance_freq",
                "min_periods",
                "future_days",
                "ic_lookback_period",
                "test_sharpe",
                "annual_return",
                "annual_volatility",
                "max_drawdown",
            ]
            table = doc.add_table(rows=1 + len(summary_df), cols=len(columns))
            table.style = "Light List Accent 1"
            header = table.rows[0].cells
            for idx, col in enumerate(columns):
                header[idx].text = col

            for row_idx, (_, row) in enumerate(summary_df.iterrows(), start=1):
                cells = table.rows[row_idx].cells
                for col_idx, col in enumerate(columns):
                    value = row[col]
                    if isinstance(value, float):
                        cells[col_idx].text = f"{value:.4f}"
                    else:
                        cells[col_idx].text = str(value)

        doc.add_heading("2. 指数详细指标", level=1)
        for code, result in results.items():
            doc.add_heading(code, level=2)
            metrics = result.best_trial.test_metrics
            paragraph = doc.add_paragraph()
            paragraph.add_run("测试集指标：\n")
            for key, value in metrics.items():
                paragraph.add_run(f"- {key}: {value:.4f}\n")
            paragraph.add_run(
                f"最优参数：{json.dumps(result.best_params, ensure_ascii=False)}"
            )

        if heatmap_path and heatmap_path.exists():
            doc.add_heading("3. 信号热力图", level=1)
            doc.add_paragraph("以下热力图展示近期信号 Z-Score 分布情况。")
            doc.add_picture(str(heatmap_path), width=Inches(6.5))

        output_path = output_dir / word_name
        doc.save(output_path)
        LOGGER.info("Saved Word report to %s", output_path)
        return output_path
